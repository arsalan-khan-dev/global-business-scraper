"""
╔══════════════════════════════════════════════════════════════════════════════╗
║          GLOBAL LOCAL BUSINESS SCRAPER & LEAD GENERATOR                      ║
║                                                                              ║
╚══════════════════════════════════════════════════════════════════════════════╝

"""
import tkinter as tk
from tkinter import ttk, messagebox, filedialog, scrolledtext
import threading
import time
import re
import csv
import os
import json
import logging
import random
from datetime import datetime
from urllib.parse import quote_plus, urljoin, urlparse

# ─── Third-party imports (with graceful fallback messages) ───────────────────
try:
    import requests
    from requests.adapters import HTTPAdapter
    from urllib3.util.retry import Retry
    REQUESTS_OK = True
except ImportError:
    REQUESTS_OK = False

try:
    from bs4 import BeautifulSoup
    BS4_OK = True
except ImportError:
    BS4_OK = False

try:
    import pandas as pd
    PANDAS_OK = True
except ImportError:
    PANDAS_OK = False

try:
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    from openpyxl.utils import get_column_letter
    OPENPYXL_OK = True
except ImportError:
    OPENPYXL_OK = False

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    DOCX_OK = True
except ImportError:
    DOCX_OK = False

# ─── Logging Setup ────────────────────────────────────────────────────────────
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s [%(levelname)s] %(message)s",
    datefmt="%H:%M:%S"
)
logger = logging.getLogger("BusinessScraper")

# ═══════════════════════════════════════════════════════════════════════════════
#  LOCATION DATA
# ═══════════════════════════════════════════════════════════════════════════════

LOCATION_DATA = {
    "Pakistan": {
        "Khyber Pakhtunkhwa": {
            "Peshawar": ["Hayatabad", "University Town", "Saddar", "Cantt", "Qissa Khwani", "Gulbahar", "Firdous", "Warsak Road"],
            "Kohat": ["Junglekhel", "Cantt", "City Center", "Lachi Road", "Hangu Road"],
            "Abbottabad": ["Mansehra Road", "Havelian", "Nawan Shehr", "Cantt"],
            "Mardan": ["Rustam", "Takht Bhai", "City Center", "Cantt"],
            "Swat": ["Mingora", "Saidu Sharif", "Matta", "Bahrain", "Kalam"],
        },
        "Punjab": {
            "Lahore": ["Gulberg", "DHA", "Model Town", "Johar Town", "Cantt", "Shadman", "Garden Town", "Bahria Town"],
            "Faisalabad": ["Peoples Colony", "Ghulam Muhammad Abad", "Cantt", "Jinnah Colony", "Madina Town"],
            "Rawalpindi": ["Saddar", "Cantt", "Bahria Town", "Dhoke Hassu", "Murree Road"],
            "Multan": ["Cantt", "Shah Rukn-e-Alam", "New Multan", "Gulgasht Colony"],
            "Gujranwala": ["Trust Colony", "Satellite Town", "Cantt", "Civil Lines"],
        },
        "Sindh": {
            "Karachi": ["Clifton", "DHA", "Gulshan-e-Iqbal", "PECHS", "Saddar", "Korangi", "Malir", "Nazimabad"],
            "Hyderabad": ["Latifabad", "Qasimabad", "City", "Cantt"],
            "Sukkur": ["Rohri", "Shikarpur Road", "Airport Road"],
        },
        "Balochistan": {
            "Quetta": ["Cantt", "Satellite Town", "Jinnah Town", "Brewery Road", "Airport Road"],
            "Gwadar": ["New Town", "Old Town", "Pasni Road"],
        },
        "Islamabad": {
            "Islamabad": ["F-6", "F-7", "F-8", "G-9", "G-10", "G-11", "I-8", "Blue Area", "E-11", "DHA"],
        },
    },
    "United States": {
        "California": {
            "Los Angeles": ["Hollywood", "Downtown", "Beverly Hills", "Santa Monica", "Koreatown", "Silver Lake"],
            "San Francisco": ["Mission", "Castro", "SOMA", "Noe Valley", "Chinatown", "Financial District"],
            "San Diego": ["Gaslamp", "North Park", "La Jolla", "Hillcrest", "Pacific Beach"],
        },
        "New York": {
            "New York City": ["Manhattan", "Brooklyn", "Queens", "Bronx", "Harlem", "Williamsburg"],
            "Buffalo": ["Downtown", "Elmwood Village", "Allentown"],
            "Albany": ["Downtown", "Center Square", "Pine Hills"],
        },
        "Texas": {
            "Houston": ["Downtown", "Midtown", "Montrose", "Heights", "The Woodlands"],
            "Dallas": ["Downtown", "Uptown", "Deep Ellum", "Oak Cliff"],
            "Austin": ["Downtown", "East Austin", "South Congress", "Domain"],
        },
    },
    "United Kingdom": {
        "England": {
            "London": ["Westminster", "Canary Wharf", "Shoreditch", "Brixton", "Camden", "Kensington"],
            "Manchester": ["City Centre", "Deansgate", "Northern Quarter", "Ancoats", "Didsbury"],
            "Birmingham": ["City Centre", "Jewellery Quarter", "Digbeth", "Moseley", "Edgbaston"],
            "Leeds": ["City Centre", "Headingley", "Chapel Allerton", "Roundhay"],
        },
        "Scotland": {
            "Edinburgh": ["Old Town", "New Town", "Leith", "Morningside"],
            "Glasgow": ["City Centre", "West End", "East End", "Southside"],
        },
        "Wales": {
            "Cardiff": ["City Centre", "Canton", "Pontcanna", "Roath", "Cathays"],
        },
    },
    "India": {
        "Maharashtra": {
            "Mumbai": ["Bandra", "Andheri", "Colaba", "Dadar", "Powai", "Juhu"],
            "Pune": ["Koregaon Park", "Kothrud", "Hadapsar", "Shivajinagar", "Wakad"],
            "Nagpur": ["Sitabuldi", "Dharampeth", "Sadar", "Civil Lines"],
        },
        "Delhi": {
            "New Delhi": ["Connaught Place", "Karol Bagh", "Lajpat Nagar", "Hauz Khas", "Saket", "Dwarka"],
        },
        "Karnataka": {
            "Bangalore": ["Indiranagar", "Koramangala", "Whitefield", "JP Nagar", "MG Road"],
            "Mysore": ["Vijayanagar", "Kuvempunagar", "Hebbal"],
        },
        "Tamil Nadu": {
            "Chennai": ["T. Nagar", "Anna Nagar", "Adyar", "Velachery", "Porur"],
        },
    },
    "UAE": {
        "Dubai": {
            "Dubai": ["Downtown Dubai", "Deira", "Bur Dubai", "JBR", "Business Bay", "DIFC", "Jumeirah"],
        },
        "Abu Dhabi": {
            "Abu Dhabi": ["Khalidiyah", "Corniche", "Al Reem Island", "Yas Island", "Mussafah"],
        },
        "Sharjah": {
            "Sharjah": ["Al Nahda", "Al Khan", "Rolla", "Muwaileh"],
        },
    },
    "Canada": {
        "Ontario": {
            "Toronto": ["Downtown", "North York", "Scarborough", "Etobicoke", "Mississauga"],
            "Ottawa": ["Downtown", "Byward Market", "Glebe", "Kanata"],
        },
        "British Columbia": {
            "Vancouver": ["Downtown", "Gastown", "Kitsilano", "Commercial Drive", "Mount Pleasant"],
            "Victoria": ["Downtown", "Fernwood", "Oak Bay"],
        },
        "Quebec": {
            "Montreal": ["Plateau", "Mile End", "Old Montreal", "Rosemont", "NDG"],
        },
    },
    "Australia": {
        "New South Wales": {
            "Sydney": ["CBD", "Surry Hills", "Newtown", "Bondi", "Manly", "Parramatta"],
            "Newcastle": ["City Centre", "Hamilton", "Islington"],
        },
        "Victoria": {
            "Melbourne": ["CBD", "Fitzroy", "Collingwood", "St Kilda", "Prahran", "Richmond"],
            "Geelong": ["City Centre", "Newtown", "Belmont"],
        },
        "Queensland": {
            "Brisbane": ["CBD", "Fortitude Valley", "South Brisbane", "New Farm"],
        },
    },
    "Germany": {
        "Bavaria": {
            "Munich": ["Altstadt", "Schwabing", "Maxvorstadt", "Haidhausen", "Neuhausen"],
            "Nuremberg": ["Altstadt", "Gostenhof", "Langwasser"],
        },
        "Berlin": {
            "Berlin": ["Mitte", "Prenzlauer Berg", "Kreuzberg", "Charlottenburg", "Neukölln"],
        },
        "Hamburg": {
            "Hamburg": ["Altstadt", "Eimsbüttel", "Altona", "Harburg", "Wandsbek"],
        },
    },
    "Saudi Arabia": {
        "Riyadh Region": {
            "Riyadh": ["Olaya", "Al Malaz", "Al Sulaimaniyah", "Diplomatic Quarter", "Al Naseem"],
        },
        "Makkah Region": {
            "Jeddah": ["Al Balad", "Al Hamra", "Al Rawdah", "Corniche", "Al Shati"],
            "Mecca": ["Ajyad", "Al Aziziyah", "Batha Quraysh"],
        },
        "Eastern Province": {
            "Dammam": ["Al Faisaliah", "Al Nakheel", "Corniche"],
            "Al Khobar": ["Downtown", "Corniche", "Al Thuqbah"],
        },
    },
}

BUSINESS_CATEGORIES = [
    "Bike Shops",
    "Car Showrooms",
    "Schools",
    "Colleges",
    "Universities",
    "Pharmacies",
    "Hospitals",
    "Restaurants",
    "Charity Organizations",
    "Electronics Shops",
    "Clothing Stores",
    "Bakeries",
    "Grocery Stores",
    "Hotels",
    "Gyms & Fitness Centers",
    "Beauty Salons",
    "Barber Shops",
    "Dentists",
    "Law Firms",
    "Accounting Firms",
    "Real Estate Agencies",
    "Banks",
    "Mobile Phone Shops",
    "Computer Repair Shops",
    "Petrol Stations",
    "Mechanics / Auto Repair",
    "Furniture Stores",
    "Jewelry Stores",
    "Bookstores",
    "Cafes & Coffee Shops",
    "Printing & Stationery",
    "Tailors / Sewing Shops",
    "Opticians / Eye Clinics",
    "Veterinary Clinics",
    "Travel Agencies",
]

# ═══════════════════════════════════════════════════════════════════════════════
#  SCRAPER ENGINE  –  v1.0  (Google + Google Maps primary sources)
# ═══════════════════════════════════════════════════════════════════════════════

USER_AGENTS = [
    "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/122.0.0.0 Safari/537.36",
    "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/121.0.0.0 Safari/537.36",
    "Mozilla/5.0 (Windows NT 10.0; Win64; x64; rv:124.0) Gecko/20100101 Firefox/124.0",
    "Mozilla/5.0 (X11; Linux x86_64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36",
    "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) AppleWebKit/605.1.15 (KHTML, like Gecko) Version/17.3 Safari/605.1.15",
    "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/123.0.0.0 Safari/537.36 Edg/123.0.0.0",
]

GOOGLE_HEADERS_POOL = [
    {
        "User-Agent": ua,
        "Accept": "text/html,application/xhtml+xml,application/xml;q=0.9,image/avif,image/webp,*/*;q=0.8",
        "Accept-Language": "en-US,en;q=0.9",
        "Accept-Encoding": "gzip, deflate, br",
        "Referer": "https://www.google.com/",
        "Connection": "keep-alive",
        "Upgrade-Insecure-Requests": "1",
        "Sec-Fetch-Dest": "document",
        "Sec-Fetch-Mode": "navigate",
        "Sec-Fetch-Site": "same-origin",
        "Cache-Control": "max-age=0",
    }
    for ua in USER_AGENTS
]


def get_session(use_google_headers=False):
    """Create a requests session with retry logic."""
    session = requests.Session()
    retry = Retry(total=3, backoff_factor=1.5, status_forcelist=[429, 500, 502, 503, 504])
    adapter = HTTPAdapter(max_retries=retry)
    session.mount("http://", adapter)
    session.mount("https://", adapter)
    if use_google_headers:
        session.headers.update(random.choice(GOOGLE_HEADERS_POOL))
    else:
        session.headers.update({
            "User-Agent": random.choice(USER_AGENTS),
            "Accept": "text/html,application/xhtml+xml,application/xml;q=0.9,*/*;q=0.8",
            "Accept-Language": "en-US,en;q=0.5",
            "Accept-Encoding": "gzip, deflate, br",
            "Connection": "keep-alive",
        })
    return session


def build_query(category, area, city, province, country):
    parts = [category, "in"]
    if area:
        parts.append(area)
    parts.append(city)
    if province and province != city:
        parts.append(province)
    parts.append(country)
    return " ".join(parts)


def extract_phone(text):
    patterns = [
        r'\+?92[\s\-]?\d{3}[\s\-]?\d{7}',
        r'0\d{3}[\s\-]?\d{7}',
        r'\+?[\d\s\-\(\)]{10,17}',
        r'\b\d{3}[\s\-]?\d{3,4}[\s\-]?\d{4}\b',
        r'\b\d{4}[\s\-]?\d{7}\b',
        r'\b\d{11}\b',
    ]
    for pattern in patterns:
        matches = re.findall(pattern, text)
        for match in matches:
            cleaned = re.sub(r'\s+', ' ', match.strip())
            digits = re.sub(r'\D', '', cleaned)
            if 7 <= len(digits) <= 15:
                return cleaned
    return ""


def extract_email(text):
    pattern = r'\b[A-Za-z0-9._%+\-]+@[A-Za-z0-9.\-]+\.[A-Za-z]{2,}\b'
    matches = re.findall(pattern, text)
    for match in matches:
        if not any(skip in match.lower() for skip in ['example', 'test', 'noreply', 'domain', 'email', 'sentry']):
            return match
    return ""


def extract_website(text, soup=None):
    if soup:
        for a in soup.find_all('a', href=True):
            href = a['href']
            if href.startswith('http') and not any(
                x in href for x in ['google', 'facebook', 'twitter', 'instagram', 'yelp', 'yellowpages', 'maps.']
            ):
                return href
    pattern = r'https?://[^\s<>"{}|\\^`\[\]]+'
    matches = re.findall(pattern, text)
    for match in matches:
        if not any(x in match for x in ['google', 'facebook', 'twitter', 'instagram', 'yelp']):
            return match.rstrip('.,)')
    return ""


def check_website_status(url):
    if not url:
        return "No"
    try:
        session = get_session()
        resp = session.head(url, timeout=6, allow_redirects=True)
        return "Yes" if resp.status_code < 400 else "No"
    except Exception:
        try:
            session = get_session()
            resp = session.get(url, timeout=8, allow_redirects=True)
            return "Yes" if resp.status_code < 400 else "No"
        except Exception:
            return "No"


# ─────────────────────────────────────────────────────────────────────────────
#  SOURCE 1 ── Google Search  (PRIMARY)
# ─────────────────────────────────────────────────────────────────────────────

def scrape_google(query, max_results=25, log_callback=None):
    """Scrape Google search results – organic listings + local pack."""
    results = []
    if not REQUESTS_OK or not BS4_OK:
        return results

    encoded_query = quote_plus(query)
    urls_to_try = [
        f"https://www.google.com/search?q={encoded_query}&num=30&hl=en&gl=us",
        f"https://www.google.com/search?q={encoded_query}&num=20&hl=en",
    ]

    for attempt, search_url in enumerate(urls_to_try):
        try:
            if log_callback:
                log_callback(f"🔍 [Google] Querying: {query}")

            session = get_session(use_google_headers=True)
            # Mimic real browser delay
            time.sleep(random.uniform(2.0, 4.5))

            resp = session.get(search_url, timeout=20)
            if resp.status_code == 429:
                if log_callback:
                    log_callback("Google rate-limited. Waiting 10s before retry…")
                time.sleep(10)
                continue

            resp.raise_for_status()
            soup = BeautifulSoup(resp.text, "lxml")

            # ── Parse organic results ─────────────────────────────────────────
            # Google wraps results in <div class="g"> or similar
            selectors = [
                "div.g",
                "div[data-sokoban-container]",
                "div.MjjYud",
                "div.tF2Cxc",
                "div.kvH3mc",
            ]
            result_divs = []
            for sel in selectors:
                result_divs = soup.select(sel)
                if result_divs:
                    break

            if log_callback:
                log_callback(f"    Google returned {len(result_divs)} organic result blocks")

            for div in result_divs[:max_results]:
                # Title
                title_tag = (div.select_one("h3") or
                             div.select_one(".DKV0Md") or
                             div.select_one(".LC20lb"))
                if not title_tag:
                    continue
                title = title_tag.get_text(strip=True)

                # URL
                link_tag = div.select_one("a[href]")
                url = ""
                if link_tag:
                    href = link_tag.get("href", "")
                    if href.startswith("/url?q="):
                        url = href.split("/url?q=")[1].split("&")[0]
                    elif href.startswith("http"):
                        url = href

                # Snippet
                snippet_tag = (div.select_one(".VwiC3b") or
                               div.select_one(".s3v9rd") or
                               div.select_one("span.aCOpRe") or
                               div.select_one(".IsZvec"))
                snippet = snippet_tag.get_text(strip=True) if snippet_tag else ""

                # Display URL
                cite_tag = div.select_one("cite")
                display_url = cite_tag.get_text(strip=True) if cite_tag else urlparse(url).netloc

                if title and len(title) > 3:
                    results.append({
                        "title": title,
                        "snippet": snippet,
                        "url": url,
                        "display_url": display_url,
                        "source": "Google",
                    })

            # ── Parse Google Local Pack (map results) ─────────────────────────
            local_selectors = [
                "div.VkpGBb",
                "div.uMdZh",
                "div[jscontroller] div.rllt__details",
                "div.cXedhc",
                "div[data-cid]",
            ]
            for sel in local_selectors:
                local_divs = soup.select(sel)
                if local_divs:
                    if log_callback:
                        log_callback(f"   📍 Found {len(local_divs)} Google local pack entries")
                    for ldiv in local_divs:
                        ltext = ldiv.get_text(" ", strip=True)
                        name_tag = ldiv.select_one("span.OSrXXb") or ldiv.select_one(".dbg0pd") or ldiv.find(["span","div"], class_=re.compile(r"name|title", re.I))
                        name = name_tag.get_text(strip=True) if name_tag else ltext[:60]
                        if name and len(name) > 3:
                            phone = extract_phone(ltext)
                            results.append({
                                "title": name,
                                "snippet": ltext[:200],
                                "url": "",
                                "display_url": "",
                                "source": "Google Local Pack",
                                "phone_hint": phone,
                            })
                    break

            if results:
                break  # Success, no need for second URL

        except requests.exceptions.ConnectionError:
            if log_callback:
                log_callback("⚠️  Network error connecting to Google")
        except requests.exceptions.Timeout:
            if log_callback:
                log_callback("⚠️  Google request timed out")
        except Exception as e:
            if log_callback:
                log_callback(f"⚠️  Google error: {type(e).__name__}: {e}")

    return results


# ─────────────────────────────────────────────────────────────────────────────
#  SOURCE 2 ── Google Maps  (BEST for local businesses)
# ─────────────────────────────────────────────────────────────────────────────

def scrape_google_maps(query, log_callback=None):
    """
    Scrape Google Maps search results.
    Uses the maps HTML endpoint which returns structured business data.
    """
    results = []
    if not REQUESTS_OK or not BS4_OK:
        return results

    encoded = quote_plus(query)
    maps_url = f"https://www.google.com/maps/search/{encoded}"

    try:
        if log_callback:
            log_callback(f" [Google Maps] Querying: {query}")

        session = get_session(use_google_headers=True)
        time.sleep(random.uniform(2.5, 5.0))

        resp = session.get(maps_url, timeout=20)
        resp.raise_for_status()

        soup = BeautifulSoup(resp.text, "lxml")
        page_text = resp.text

        # ── Strategy A: Parse embedded JSON data ──────────────────────────────
        # Google Maps embeds business data as JSON arrays in the page source
        json_blocks_found = 0

        # Pattern: look for business name arrays in the JS
        # Google Maps uses window.APP_INITIALIZATION_STATE or similar
        patterns_to_try = [
            r'"([^"]{3,80})",[^,]*"([^"]*(?:street|road|avenue|lane|plaza|colony|block|sector)[^"]*)"',
            r'\["([^"]{5,60})",[^,\[]*\[null,null,(-?\d+\.\d+),(-?\d+\.\d+)\]',
        ]

        # ── Strategy B: Parse visible HTML listings ───────────────────────────
        # When JS is disabled, Maps returns simplified HTML
        listing_containers = (
            soup.select("div[class*='section-result']") or
            soup.select("div[aria-label]") or
            soup.select("div[data-result-index]") or
            soup.select(".section-result-content")
        )

        if listing_containers:
            if log_callback:
                log_callback(f"   📍 Maps HTML mode: {len(listing_containers)} listings found")
            for item in listing_containers[:25]:
                text = item.get_text(" ", strip=True)
                name_tag = item.find(["h3", "h2", "span"], attrs={"aria-label": True}) or item.find("h3") or item.find("h2")
                name = ""
                if name_tag:
                    name = name_tag.get("aria-label", "") or name_tag.get_text(strip=True)
                if not name:
                    name = text[:70]
                if name and len(name) > 3:
                    phone = extract_phone(text)
                    results.append({
                        "title": name.strip(),
                        "snippet": text[:200],
                        "url": "",
                        "display_url": "maps.google.com",
                        "source": "Google Maps",
                        "phone_hint": phone,
                        "maps_link": maps_url,
                    })

        # ── Strategy C: Extract from JS data blobs in <script> tags ──────────
        if not results:
            scripts = soup.find_all("script")
            for script in scripts:
                src = script.string or ""
                # Look for arrays that look like business names + coordinates
                # Pattern: ["Business Name", null, [null, lat, lng]]
                name_coord_matches = re.findall(
                    r'\["([A-Za-z][^"]{4,70})",\s*null,\s*\[(?:null,\s*)?(-?\d{1,3}\.\d{4,}),\s*(-?\d{1,3}\.\d{4,})\]',
                    src
                )
                if name_coord_matches:
                    for match in name_coord_matches[:25]:
                        name, lat, lng = match
                        # filter obvious garbage
                        if any(c in name for c in ['\\', '/', '{', '}']):
                            continue
                        maps_link_item = f"https://www.google.com/maps/search/{quote_plus(name + ' ' + query.split('in')[-1].strip())}"
                        results.append({
                            "title": name,
                            "snippet": f"Lat: {lat}, Lng: {lng}",
                            "url": "",
                            "display_url": "maps.google.com",
                            "source": "Google Maps JS",
                            "phone_hint": "",
                            "maps_link": maps_link_item,
                            "lat": lat,
                            "lng": lng,
                        })
                    if results:
                        if log_callback:
                            log_callback(f"   📍 Maps JS data: extracted {len(results)} businesses")
                        break

        if not results and log_callback:
            log_callback("   ℹ️  Google Maps returned limited HTML (JS rendering needed)")

    except requests.exceptions.ConnectionError:
        if log_callback:
            log_callback("⚠️  Network error connecting to Google Maps")
    except Exception as e:
        if log_callback:
            log_callback(f"⚠️  Google Maps error: {type(e).__name__}: {e}")

    return results


# ─────────────────────────────────────────────────────────────────────────────
#  SOURCE 3 ── Google Maps API
# ─────────────────────────────────────────────────────────────────────────────

def scrape_google_maps_places(query, city, country, log_callback=None):
    """
    Use Google's public search endpoint that powers the Maps 'local results'
    sidebar. No API key required for basic usage.
    Returns structured business list.
    """
    results = []
    if not REQUESTS_OK or not BS4_OK:
        return results

    # This hits Google's local knowledge graph panel
    search_term = quote_plus(query)
    urls = [
        f"https://www.google.com/search?q={search_term}&tbm=lcl&hl=en",   # Local tab
        f"https://www.google.com/search?q={search_term}+site:maps.google.com&hl=en",
    ]

    for url in urls:
        try:
            if log_callback:
                log_callback(f"🔍 [Google Local] {query}")

            session = get_session(use_google_headers=True)
            time.sleep(random.uniform(2.0, 4.0))

            resp = session.get(url, timeout=20)
            if resp.status_code == 429:
                if log_callback:
                    log_callback(" Google rate limit. Pausing 15s…")
                time.sleep(15)
                continue

            resp.raise_for_status()
            soup = BeautifulSoup(resp.text, "lxml")

            # Google Local Results (tbm=lcl) selectors
            # These appear as cards with business name, rating, address
            local_result_selectors = [
                "div.rllt__details",       # Standard local result detail
                "div[class*='local']",
                "div.uMdZh",
                "div.VkpGBb",
                "div.cXedhc",
                "div[data-hveid] h3",
                ".rllt__wrapped",
                "div[jscontroller] .dbg0pd",
            ]

            found_items = []
            for sel in local_result_selectors:
                found_items = soup.select(sel)
                if found_items:
                    break

            if log_callback:
                log_callback(f"   ✅ Google Local: {len(found_items)} items found")

            for item in found_items[:20]:
                item_text = item.get_text(" ", strip=True)

                # Try to get the business name specifically
                name_tag = (item.select_one(".OSrXXb") or
                            item.select_one(".dbg0pd") or
                            item.select_one("span[role='heading']") or
                            item.find("span", class_=re.compile(r"name|title|heading", re.I)))

                name = name_tag.get_text(strip=True) if name_tag else item_text[:60]
                name = name.strip()

                if not name or len(name) < 3:
                    continue

                phone = extract_phone(item_text)
                addr_match = re.search(r'(\d+[^,\n]{5,60}(?:Street|St|Road|Rd|Avenue|Ave|Lane|Ln|Colony|Block|Sector|Plaza|Town)[^,\n]*)', item_text, re.I)
                address = addr_match.group(1).strip() if addr_match else ""

                maps_link = f"https://www.google.com/maps/search/{quote_plus(name + ' ' + city + ' ' + country)}"

                results.append({
                    "title": name,
                    "snippet": item_text[:200],
                    "url": "",
                    "display_url": "google.com/maps",
                    "source": "Google Local",
                    "phone_hint": phone,
                    "address_hint": address,
                    "maps_link": maps_link,
                })

            if results:
                break

        except Exception as e:
            if log_callback:
                log_callback(f"⚠️  Google Local error: {e}")

    return results


# ─────────────────────────────────────────────────────────────────────────────
#  SOURCE 4 ── DuckDuckGo  (FALLBACK)
# ─────────────────────────────────────────────────────────────────────────────

def scrape_duckduckgo(query, max_results=20, log_callback=None):
    results = []
    if not REQUESTS_OK or not BS4_OK:
        return results

    session = get_session()
    search_url = f"https://html.duckduckgo.com/html/?q={quote_plus(query)}"

    try:
        if log_callback:
            log_callback(f"🔍 [DDG] Querying: {query}")
        time.sleep(random.uniform(1.5, 3.0))
        resp = session.get(search_url, timeout=15)
        resp.raise_for_status()
        soup = BeautifulSoup(resp.text, "lxml")

        result_divs = (soup.select(".result__body") or
                       soup.select(".results_links") or
                       soup.find_all("div", class_=re.compile(r"result")))

        if log_callback:
            log_callback(f"   DDG: {len(result_divs)} blocks found")

        for div in result_divs[:max_results]:
            title_tag = (div.select_one(".result__title a") or
                         div.select_one("a.result__url") or div.find("a"))
            snippet_tag = (div.select_one(".result__snippet") or
                           div.select_one(".result__body"))
            url_tag = div.select_one(".result__url") or div.select_one(".result__extras__url")

            if not title_tag:
                continue

            title = title_tag.get_text(strip=True)
            snippet = snippet_tag.get_text(strip=True) if snippet_tag else ""
            url = title_tag.get("href", "")
            display_url = url_tag.get_text(strip=True) if url_tag else ""

            if title and len(title) > 3:
                results.append({
                    "title": title,
                    "snippet": snippet,
                    "url": url,
                    "display_url": display_url,
                    "source": "DuckDuckGo",
                })

    except requests.exceptions.ConnectionError:
        if log_callback:
            log_callback("⚠️  Network error – DuckDuckGo")
    except Exception as e:
        if log_callback:
            log_callback(f"⚠️  DDG error: {e}")

    return results


# ─────────────────────────────────────────────────────────────────────────────
#  SOURCE 5 ── Bing  (FALLBACK)
# ─────────────────────────────────────────────────────────────────────────────

def scrape_bing(query, max_results=20, log_callback=None):
    results = []
    if not REQUESTS_OK or not BS4_OK:
        return results

    session = get_session()
    search_url = f"https://www.bing.com/search?q={quote_plus(query)}&count=30"

    try:
        if log_callback:
            log_callback(f"🔍 [Bing] Querying: {query}")
        time.sleep(random.uniform(2.0, 4.0))
        resp = session.get(search_url, timeout=15)
        resp.raise_for_status()
        soup = BeautifulSoup(resp.text, "lxml")

        result_items = soup.select("li.b_algo") or soup.select(".b_results > li")
        if log_callback:
            log_callback(f"   Bing: {len(result_items)} items")

        for item in result_items[:max_results]:
            title_tag = item.select_one("h2 a") or item.select_one("h3 a")
            snippet_tag = item.select_one(".b_caption p") or item.select_one("p")

            if not title_tag:
                continue

            title = title_tag.get_text(strip=True)
            url = title_tag.get("href", "")
            snippet = snippet_tag.get_text(strip=True) if snippet_tag else ""

            if title and url.startswith("http"):
                results.append({
                    "title": title,
                    "snippet": snippet,
                    "url": url,
                    "display_url": urlparse(url).netloc,
                    "source": "Bing",
                })

    except Exception as e:
        if log_callback:
            log_callback(f"⚠️  Bing error: {e}")

    return results


# ─────────────────────────────────────────────────────────────────────────────
#  Business page enrichment
# ─────────────────────────────────────────────────────────────────────────────

def scrape_business_page(url, category, location_str, log_callback=None):
    if not url or not url.startswith("http"):
        return None

    skip_domains = [
        "google.com", "facebook.com", "twitter.com", "instagram.com",
        "youtube.com", "wikipedia.org", "linkedin.com", "amazon.com",
        "yelp.com", "tripadvisor.com", "yellowpages.com", "maps.google",
    ]
    domain = urlparse(url).netloc.lower()
    if any(skip in domain for skip in skip_domains):
        return None

    try:
        session = get_session()
        time.sleep(random.uniform(1.0, 2.5))
        resp = session.get(url, timeout=12)
        resp.raise_for_status()
        soup = BeautifulSoup(resp.text, "lxml")

        for tag in soup(["script", "style", "noscript"]):
            tag.decompose()

        page_text = soup.get_text(separator=" ", strip=True)
        phone = extract_phone(page_text)
        email = extract_email(page_text)
        website = url

        address = ""
        address_patterns = [
            r'\d+\s+[\w\s]+(?:Street|St|Avenue|Ave|Road|Rd|Lane|Ln|Drive|Dr|Boulevard|Blvd|Plaza|Colony|Town|Sector|Block|Phase|Plot)[,\s\w]*',
            r'(?:Address|Location|Located at|Find us at)[:\s]+([^\n\.]{10,120})',
        ]
        for pat in address_patterns:
            m = re.search(pat, page_text, re.IGNORECASE)
            if m:
                address = m.group(0)[:120].strip()
                break

        if not address:
            address = location_str

        name = soup.find("title")
        name = (name.get_text(strip=True).split("|")[0].split("-")[0].strip()
                if name else urlparse(url).netloc)

        maps_link = f"https://www.google.com/maps/search/{quote_plus(name + ' ' + location_str)}"

        return {
            "name": name[:80],
            "address": address[:120],
            "phone": phone[:30] if phone else "",
            "email": email[:60] if email else "",
            "website": website,
            "maps_link": maps_link,
        }

    except Exception:
        return None


def parse_search_result_to_business(result, category, location_str, city="", country=""):
    title = result.get("title", "").strip()
    snippet = result.get("snippet", "").strip()
    url = result.get("url", "")
    source = result.get("source", "")

    combined_text = f"{title} {snippet}"

    phone = result.get("phone_hint", "") or extract_phone(combined_text)
    email = extract_email(combined_text)

    # Address
    address = result.get("address_hint", "")
    if not address:
        addr_patterns = [
            r'\d+[\w\s,\-]+(?:Street|St|Avenue|Ave|Road|Rd|Lane|Colony|Block|Sector|Town|Plaza)[,\s\w]*',
            r'(?:located|address|find us)[:\s]+([^\.\n]{10,80})',
        ]
        for pat in addr_patterns:
            m = re.search(pat, combined_text, re.IGNORECASE)
            if m:
                address = m.group(0)[:120].strip()
                break

    if not address:
        address = location_str

    maps_link = result.get("maps_link") or \
        f"https://www.google.com/maps/search/{quote_plus(title + ' ' + location_str)}"

    return {
        "name": title[:80],
        "category": category,
        "address": address[:120],
        "phone": phone[:30] if phone else "",
        "email": email[:60] if email else "",
        "website": url if url.startswith("http") else "",
        "website_status": "",
        "maps_link": maps_link,
        "source": source,
    }


# ─────────────────────────────────────────────────────────────────────────────
#  MAIN SCRAPER ORCHESTRATOR
# ─────────────────────────────────────────────────────────────────────────────

def run_scraper(country, province, city, area, category, max_results,
                log_callback, progress_callback, result_callback, done_callback):
    """
    Main scraping orchestrator. Priority order:
      1. Google Local (tbm=lcl)
      2. Google Maps search
      3. Google organic search
      4. DuckDuckGo (fallback)
      5. Bing (fallback)
    """
    businesses = []
    seen_names = set()
    location_str = f"{area}, {city}, {province}, {country}".strip(", ").strip(", ")

    base_query  = build_query(category, area, city, province, country)
    short_query = f"{category} in {city} {country}"
    maps_query  = f"{category} near {city}"

    log_callback("=" * 50)
    log_callback("  NEW SEARCH STARTED")
    log_callback("=" * 50)
    log_callback(f"📋 Location: {location_str}")
    log_callback(f"📦 Category: {category}")
    log_callback(f"🔢 Max results: {max_results}")
    log_callback("─" * 50)

    total_raw = []

    # ── Phase 1: Google Local (best for businesses) ────────────────────────
    progress_callback(5)
    log_callback("🔎 Phase 1: Searching Google Local results…")
    raw = scrape_google_maps_places(base_query, city, country, log_callback=log_callback)
    total_raw.extend(raw)
    if not raw:
        raw2 = scrape_google_maps_places(short_query, city, country, log_callback=log_callback)
        total_raw.extend(raw2)

    # ── Phase 2: Google Maps ────────────────────────────────────────────────
    progress_callback(15)
    log_callback("🗺️  Phase 2: Searching Google Maps…")
    maps_raw = scrape_google_maps(base_query, log_callback=log_callback)
    total_raw.extend(maps_raw)
    time.sleep(random.uniform(2.0, 3.5))

    if not maps_raw:
        maps_raw2 = scrape_google_maps(short_query, log_callback=log_callback)
        total_raw.extend(maps_raw2)

    # ── Phase 3: Google organic search ────────────────────────────────────
    progress_callback(28)
    log_callback("🔍 Phase 3: Searching Google organic results…")
    google_queries = [base_query, short_query]
    if area:
        google_queries.append(f"{category} near {area} {city}")

    for i, q in enumerate(google_queries):
        if len(total_raw) >= max_results * 3:
            break
        graw = scrape_google(q, max_results=20, log_callback=log_callback)
        total_raw.extend(graw)
        time.sleep(random.uniform(3.0, 6.0))   # be polite to Google

    # ── Phase 4: DuckDuckGo fallback ──────────────────────────────────────
    progress_callback(45)
    log_callback("🔍 Phase 4: DuckDuckGo fallback…")
    ddg_raw = scrape_duckduckgo(base_query, max_results=20, log_callback=log_callback)
    total_raw.extend(ddg_raw)
    time.sleep(random.uniform(1.5, 3.0))

    # ── Phase 5: Bing fallback ─────────────────────────────────────────────
    progress_callback(52)
    log_callback("🔍 Phase 5: Bing fallback…")
    bing_raw = scrape_bing(base_query, max_results=20, log_callback=log_callback)
    total_raw.extend(bing_raw)

    log_callback(f"✅ Total raw results collected: {len(total_raw)}")
    log_callback("─" * 50)

    # ── Phase 6: Parse & deduplicate ──────────────────────────────────────
    progress_callback(58)
    log_callback("🔄 Processing and deduplicating results…")

    # Sort: Google Local/Maps first, then organic, then fallbacks
    source_priority = {"Google Local": 0, "Google Maps": 1, "Google Maps JS": 2,
                       "Google": 3, "DuckDuckGo": 4, "Bing": 5}
    total_raw.sort(key=lambda r: source_priority.get(r.get("source", ""), 9))

    for result in total_raw:
        if len(businesses) >= max_results:
            break

        title = result.get("title", "").strip()
        if not title or len(title) < 4:
            continue
        # Skip obviously non-business titles
        skip_keywords = ["wikipedia", "how to", "what is", "list of", "top 10",
                         "best ", "reviews", "contact us", "home - ", "welcome to"]
        if any(kw in title.lower() for kw in skip_keywords):
            continue

        norm = re.sub(r'\W+', '', title.lower())
        if norm in seen_names:
            continue
        seen_names.add(norm)

        biz = parse_search_result_to_business(result, category, location_str, city, country)
        if biz and biz["name"]:
            businesses.append(biz)
            result_callback(biz)

    log_callback(f"✅ Unique businesses after dedup: {len(businesses)}")

    # ── Phase 7: Enrich top results ────────────────────────────────────────
    progress_callback(65)
    enrich_limit = min(12, len(businesses))
    log_callback(f"🌐 Enriching top {enrich_limit} results by visiting pages…")

    for idx, biz in enumerate(businesses[:enrich_limit]):
        progress_callback(65 + int(idx / max(enrich_limit, 1) * 22))
        url = biz.get("website", "")
        if not url:
            continue
        log_callback(f"   ↳ {url[:65]}…")
        enriched = scrape_business_page(url, category, location_str, log_callback)
        if enriched:
            if not biz["phone"] and enriched.get("phone"):
                biz["phone"] = enriched["phone"]
            if not biz["email"] and enriched.get("email"):
                biz["email"] = enriched["email"]
            if not biz["address"] or biz["address"] == location_str:
                if enriched.get("address") and enriched["address"] != location_str:
                    biz["address"] = enriched["address"]

    # ── Phase 8: Check website statuses ───────────────────────────────────
    progress_callback(90)
    log_callback("🔗 Checking website statuses…")
    for biz in businesses:
        biz["website_status"] = check_website_status(biz.get("website", "")) if biz.get("website") else "No"

    # ── Done ───────────────────────────────────────────────────────────────
    progress_callback(100)
    log_callback("─" * 50)
    log_callback(f"🎉 Scraping complete! Found {len(businesses)} businesses.")
    log_callback(f"📊 Sources used: Google Local, Google Maps, Google, DDG, Bing")
    done_callback(businesses)


# ═══════════════════════════════════════════════════════════════════════════════
#  EXPORT FUNCTIONS
# ═══════════════════════════════════════════════════════════════════════════════

COLUMNS = ["name", "category", "address", "phone", "email", "website", "website_status", "maps_link", "source"]
HEADERS = ["Business Name", "Category", "Address", "Phone", "Email", "Website", "Website Status", "Google Maps", "Source"]


def export_csv(businesses, filepath):
    with open(filepath, "w", newline="", encoding="utf-8-sig") as f:
        writer = csv.DictWriter(f, fieldnames=COLUMNS, extrasaction="ignore")
        writer.writeheader()
        writer.writerows(businesses)
    return True


def export_excel(businesses, filepath, search_info):
    if not OPENPYXL_OK:
        raise ImportError("openpyxl not installed. Run: pip install openpyxl")

    wb = Workbook()
    ws = wb.active
    ws.title = "Business Leads"

    ws.merge_cells("A1:I1")
    title_cell = ws["A1"]
    title_cell.value = "🌍 Global Business Scraper – Lead Report (Google Powered)"
    title_cell.font = Font(name="Calibri", bold=True, size=16, color="FFFFFF")
    title_cell.fill = PatternFill(start_color="1a237e", end_color="1a237e", fill_type="solid")
    title_cell.alignment = Alignment(horizontal="center", vertical="center")
    ws.row_dimensions[1].height = 36

    meta_rows = [
        ("Generated:", datetime.now().strftime("%Y-%m-%d %H:%M")),
        ("Location:", search_info.get("location", "")),
        ("Category:", search_info.get("category", "")),
        ("Total Businesses:", str(len(businesses))),
    ]
    for r_idx, (label, value) in enumerate(meta_rows, start=2):
        ws.cell(row=r_idx, column=1, value=label).font = Font(bold=True, color="1a237e")
        ws.cell(row=r_idx, column=2, value=value)

    header_row = 7
    header_fill = PatternFill(start_color="283593", end_color="283593", fill_type="solid")
    header_font = Font(name="Calibri", bold=True, color="FFFFFF", size=11)
    thin = Side(style="thin", color="CCCCCC")
    border = Border(left=thin, right=thin, top=thin, bottom=thin)

    for col_idx, header in enumerate(HEADERS, start=1):
        cell = ws.cell(row=header_row, column=col_idx, value=header)
        cell.font = header_font
        cell.fill = header_fill
        cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
        cell.border = border
    ws.row_dimensions[header_row].height = 28

    alt_fill    = PatternFill(start_color="E8EAF6", end_color="E8EAF6", fill_type="solid")
    normal_fill = PatternFill(start_color="FFFFFF", end_color="FFFFFF", fill_type="solid")

    for row_idx, biz in enumerate(businesses, start=header_row + 1):
        fill = alt_fill if row_idx % 2 == 0 else normal_fill
        values = [biz.get(c, "") for c in COLUMNS]
        for col_idx, value in enumerate(values, start=1):
            cell = ws.cell(row=row_idx, column=col_idx, value=value)
            cell.fill = fill
            cell.border = border
            cell.alignment = Alignment(wrap_text=True, vertical="center")

    col_widths = [35, 20, 40, 18, 30, 40, 14, 35, 16]
    for i, width in enumerate(col_widths, start=1):
        ws.column_dimensions[get_column_letter(i)].width = width

    ws.freeze_panes = ws.cell(row=header_row + 1, column=1)
    wb.save(filepath)
    return True


def export_word(businesses, filepath, search_info):
    if not DOCX_OK:
        raise ImportError("python-docx not installed. Run: pip install python-docx")

    doc = Document()
    from docx.shared import Cm
    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    title = doc.add_heading("Global Business Scraper – Lead Report", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.color.rgb = RGBColor(0x1a, 0x23, 0x7e)

    doc.add_paragraph()

    meta_table = doc.add_table(rows=4, cols=2)
    meta_table.style = "Table Grid"
    meta_data = [
        ("Generated", datetime.now().strftime("%Y-%m-%d %H:%M")),
        ("Location", search_info.get("location", "N/A")),
        ("Category", search_info.get("category", "N/A")),
        ("Total Businesses Found", str(len(businesses))),
    ]
    for i, (label, value) in enumerate(meta_data):
        meta_table.cell(i, 0).text = label
        meta_table.cell(i, 1).text = value
        meta_table.cell(i, 0).paragraphs[0].runs[0].bold = True

    doc.add_paragraph()
    doc.add_heading("Business Listings", level=1)

    table = doc.add_table(rows=1, cols=len(HEADERS))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    hdr_cells = table.rows[0].cells
    for i, hdr in enumerate(HEADERS):
        hdr_cells[i].text = hdr
        hdr_cells[i].paragraphs[0].runs[0].bold = True

    for biz in businesses:
        row_cells = table.add_row().cells
        values = [biz.get(c, "") for c in COLUMNS]
        for i, val in enumerate(values):
            row_cells[i].text = str(val)

    doc.add_paragraph()
    footer = doc.add_paragraph(
        f"Report generated by Global Business Scraper on {datetime.now().strftime('%Y-%m-%d')}"
    )
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.runs[0].font.size = Pt(9)
    footer.runs[0].font.color.rgb = RGBColor(0x90, 0x90, 0x90)

    doc.save(filepath)
    return True


# ═══════════════════════════════════════════════════════════════════════════════
#  MAIN GUI APPLICATION
# ═══════════════════════════════════════════════════════════════════════════════

class BusinessScraperApp(tk.Tk):
    C_BG        = "#F0F2F5"
    C_SIDEBAR   = "#1a237e"
    C_ACCENT    = "#283593"
    C_BTN       = "#3949ab"
    C_BTN_HOV   = "#283593"
    C_SUCCESS   = "#2e7d32"
    C_WARNING   = "#e65100"
    C_TEXT      = "#212121"
    C_TEXT_LITE = "#FFFFFF"
    C_ROW_ALT   = "#E8EAF6"
    C_HEADER_BG = "#283593"

    def __init__(self):
        super().__init__()
        self.title("🌍 Global Business Scraper & Lead Generator  [Google Powered v2.0]")
        self.geometry("1440x920")
        self.minsize(1100, 700)
        self.configure(bg=self.C_BG)

        self.businesses: list[dict] = []
        self.scrape_thread = None
        self.is_scraping = False

        self._setup_styles()
        self._build_ui()
        self._check_dependencies()

    def _setup_styles(self):
        style = ttk.Style(self)
        style.theme_use("clam")
        style.configure("Business.Treeview",
            background="white", foreground=self.C_TEXT,
            rowheight=26, fieldbackground="white", font=("Segoe UI", 9))
        style.configure("Business.Treeview.Heading",
            background=self.C_HEADER_BG, foreground="white",
            font=("Segoe UI", 9, "bold"), relief="flat")
        style.map("Business.Treeview.Heading",
            background=[("active", self.C_ACCENT)])
        style.map("Business.Treeview",
            background=[("selected", self.C_ACCENT)],
            foreground=[("selected", "white")])
        style.configure("Green.Horizontal.TProgressbar",
            troughcolor="#E0E0E0", background=self.C_SUCCESS,
            lightcolor=self.C_SUCCESS, darkcolor=self.C_SUCCESS)

    def _build_ui(self):
        banner = tk.Frame(self, bg=self.C_SIDEBAR, height=70)
        banner.pack(fill="x", side="top")
        banner.pack_propagate(False)

        tk.Label(banner,
            text="🌍  Global Business Scraper  •  Google Powered v2.0",
            font=("Segoe UI", 19, "bold"),
            bg=self.C_SIDEBAR, fg="white").pack(side="left", padx=24, pady=16)

        self.status_label = tk.Label(banner, text="● Ready",
            font=("Segoe UI", 10), bg=self.C_SIDEBAR, fg="#a5d6a7")
        self.status_label.pack(side="right", padx=20)

        body = tk.Frame(self, bg=self.C_BG)
        body.pack(fill="both", expand=True, padx=10, pady=8)

        left = tk.Frame(body, bg="white", width=295,
            relief="flat", bd=0, highlightbackground="#BDBDBD", highlightthickness=1)
        left.pack(side="left", fill="y", padx=(0, 8))
        left.pack_propagate(False)
        self._build_left_panel(left)

        right = tk.Frame(body, bg=self.C_BG)
        right.pack(side="left", fill="both", expand=True)
        self._build_right_panel(right)

    def _build_left_panel(self, parent):
        canvas = tk.Canvas(parent, bg="white", bd=0, highlightthickness=0)
        scrollbar = ttk.Scrollbar(parent, orient="vertical", command=canvas.yview)
        scroll_frame = tk.Frame(canvas, bg="white")
        scroll_frame.bind("<Configure>",
            lambda e: canvas.configure(scrollregion=canvas.bbox("all")))
        canvas.create_window((0, 0), window=scroll_frame, anchor="nw")
        canvas.configure(yscrollcommand=scrollbar.set)
        canvas.pack(side="left", fill="both", expand=True)
        scrollbar.pack(side="right", fill="y")

        def _section(text):
            tk.Label(scroll_frame, text=text, font=("Segoe UI", 9, "bold"),
                bg=self.C_SIDEBAR, fg="white", padx=8, pady=4,
                anchor="w").pack(fill="x", pady=(10, 2))

        def _label(text):
            tk.Label(scroll_frame, text=text, font=("Segoe UI", 9),
                bg="white", fg="#424242", anchor="w").pack(fill="x", padx=12, pady=(4, 0))

        def _combo(var, values, state="readonly", width=28):
            cb = ttk.Combobox(scroll_frame, textvariable=var, values=values,
                state=state, width=width, font=("Segoe UI", 9))
            cb.pack(padx=12, pady=(0, 2), fill="x")
            return cb

        tk.Label(scroll_frame, text="🔎  Search Filters",
            font=("Segoe UI", 13, "bold"),
            bg="white", fg=self.C_SIDEBAR).pack(pady=(16, 4))
        ttk.Separator(scroll_frame).pack(fill="x", padx=12)

        _section("  📍 Location")
        _label("Country *")
        self.v_country = tk.StringVar()
        self.cb_country = _combo(self.v_country, sorted(LOCATION_DATA.keys()))
        self.cb_country.bind("<<ComboboxSelected>>", self._on_country_change)

        _label("Province / State *")
        self.v_province = tk.StringVar()
        self.cb_province = _combo(self.v_province, [])
        self.cb_province.bind("<<ComboboxSelected>>", self._on_province_change)

        _label("City *")
        self.v_city = tk.StringVar()
        self.cb_city = _combo(self.v_city, [])
        self.cb_city.bind("<<ComboboxSelected>>", self._on_city_change)

        _label("Area / Locality")
        self.v_area = tk.StringVar()
        self.cb_area = _combo(self.v_area, [], state="normal")

        _section("  🏢 Business Category")
        _label("Select Category *")
        self.v_category = tk.StringVar()
        self.cb_category = _combo(self.v_category, BUSINESS_CATEGORIES)

        _label("Custom Category (optional)")
        self.v_custom_cat = tk.StringVar()
        tk.Entry(scroll_frame, textvariable=self.v_custom_cat,
            font=("Segoe UI", 9), bd=1, relief="solid").pack(padx=12, pady=(0, 2), fill="x")

        _section("  ⚙️  Options")
        _label("Max Results")
        self.v_limit = tk.StringVar(value="50")
        _combo(self.v_limit, ["25", "50", "100", "200", "500"])

        # Source info label
        tk.Label(scroll_frame,
            text="🔎 Sources: Google Local,\n   Google Maps, DDG, Bing",
            font=("Segoe UI", 8), bg="white", fg="#1a237e",
            justify="left").pack(padx=12, pady=(6, 0), anchor="w")

        tk.Frame(scroll_frame, height=8, bg="white").pack()
        self.btn_search = tk.Button(scroll_frame, text="🔍  Start Scraping",
            font=("Segoe UI", 11, "bold"),
            bg=self.C_BTN, fg="white", bd=0, padx=12, pady=10,
            cursor="hand2", activebackground=self.C_BTN_HOV, activeforeground="white",
            command=self._start_scraping)
        self.btn_search.pack(padx=12, fill="x", pady=(4, 2))

        self.btn_stop = tk.Button(scroll_frame, text="⛔  Stop",
            font=("Segoe UI", 10), bg="#c62828", fg="white", bd=0,
            padx=12, pady=8, cursor="hand2",
            activebackground="#b71c1c", activeforeground="white",
            command=self._stop_scraping, state="disabled")
        self.btn_stop.pack(padx=12, fill="x", pady=(0, 4))

        _section("  💾 Export Results")
        for label, cmd in [
            ("📄  Export CSV",   self._export_csv),
            ("📊  Export Excel", self._export_excel),
            ("📝  Export Word",  self._export_word),
        ]:
            tk.Button(scroll_frame, text=label,
                font=("Segoe UI", 9), bg="#1565C0", fg="white", bd=0,
                padx=8, pady=6, cursor="hand2",
                activebackground="#0D47A1", activeforeground="white",
                command=cmd).pack(padx=12, fill="x", pady=2)

        tk.Button(scroll_frame, text="🗑️  Clear Results",
            font=("Segoe UI", 9), bg="#757575", fg="white", bd=0,
            padx=8, pady=6, cursor="hand2",
            activebackground="#616161", activeforeground="white",
            command=self._clear_results).pack(padx=12, fill="x", pady=(2, 12))

    def _build_right_panel(self, parent):
        stats_bar = tk.Frame(parent, bg="white", height=42,
            highlightbackground="#BDBDBD", highlightthickness=1)
        stats_bar.pack(fill="x", pady=(0, 6))
        stats_bar.pack_propagate(False)

        self.lbl_count = tk.Label(stats_bar,
            text="  Total Found: 0  |  Showing: 0",
            font=("Segoe UI", 10, "bold"),
            bg="white", fg=self.C_SIDEBAR)
        self.lbl_count.pack(side="left", padx=16, pady=8)

        self.lbl_location_display = tk.Label(stats_bar, text="",
            font=("Segoe UI", 9), bg="white", fg="#616161")
        self.lbl_location_display.pack(side="right", padx=16)

        nb = ttk.Notebook(parent)
        nb.pack(fill="both", expand=True)

        results_tab = tk.Frame(nb, bg=self.C_BG)
        nb.add(results_tab, text="  📋 Results  ")
        self._build_results_tab(results_tab)

        log_tab = tk.Frame(nb, bg=self.C_BG)
        nb.add(log_tab, text="  📜 Activity Log  ")
        self._build_log_tab(log_tab)

        prog_frame = tk.Frame(parent, bg=self.C_BG)
        prog_frame.pack(fill="x", pady=(6, 0))

        self.progress_var = tk.IntVar(value=0)
        self.progress_bar = ttk.Progressbar(prog_frame, variable=self.progress_var,
            maximum=100, style="Green.Horizontal.TProgressbar", length=400)
        self.progress_bar.pack(side="left", fill="x", expand=True, padx=(0, 10))

        self.lbl_progress = tk.Label(prog_frame, text="0%",
            font=("Segoe UI", 9, "bold"), bg=self.C_BG, fg=self.C_SUCCESS, width=5)
        self.lbl_progress.pack(side="right")

    def _build_results_tab(self, parent):
        filter_frame = tk.Frame(parent, bg=self.C_BG)
        filter_frame.pack(fill="x", pady=(4, 4))

        tk.Label(filter_frame, text="🔎 Filter:", font=("Segoe UI", 9),
            bg=self.C_BG).pack(side="left", padx=(0, 4))
        self.v_filter = tk.StringVar()
        self.v_filter.trace("w", self._apply_filter)
        tk.Entry(filter_frame, textvariable=self.v_filter,
            font=("Segoe UI", 9), width=30, bd=1, relief="solid").pack(side="left")

        tree_frame = tk.Frame(parent, bg=self.C_BG)
        tree_frame.pack(fill="both", expand=True)

        columns = ("name", "category", "address", "phone", "email", "website", "status", "source", "maps")
        self.tree = ttk.Treeview(tree_frame, columns=columns,
            show="headings", style="Business.Treeview")

        col_cfg = [
            ("name",     "Business Name",  190, "w"),
            ("category", "Category",       110, "center"),
            ("address",  "Address",        200, "w"),
            ("phone",    "Phone",          110, "center"),
            ("email",    "Email",          170, "w"),
            ("website",  "Website",        160, "w"),
            ("status",   "Live?",           60, "center"),
            ("source",   "Source",         100, "center"),
            ("maps",     "Maps",            70, "center"),
        ]
        for cid, cname, width, anchor in col_cfg:
            self.tree.heading(cid, text=cname,
                command=lambda c=cid: self._sort_column(c))
            self.tree.column(cid, width=width, anchor=anchor, minwidth=50)

        self.tree.tag_configure("even",        background=self.C_ROW_ALT)
        self.tree.tag_configure("odd",         background="white")
        self.tree.tag_configure("has_contact", foreground="#1b5e20")
        self.tree.tag_configure("google",      foreground="#1a237e")

        vsb = ttk.Scrollbar(tree_frame, orient="vertical",   command=self.tree.yview)
        hsb = ttk.Scrollbar(tree_frame, orient="horizontal", command=self.tree.xview)
        self.tree.configure(yscrollcommand=vsb.set, xscrollcommand=hsb.set)
        self.tree.grid(row=0, column=0, sticky="nsew")
        vsb.grid(row=0, column=1, sticky="ns")
        hsb.grid(row=1, column=0, sticky="ew")
        tree_frame.rowconfigure(0, weight=1)
        tree_frame.columnconfigure(0, weight=1)

        self.context_menu = tk.Menu(self, tearoff=0)
        self.context_menu.add_command(label="📋 Copy Business Name", command=self._copy_name)
        self.context_menu.add_command(label="📞 Copy Phone",         command=self._copy_phone)
        self.context_menu.add_command(label="✉️  Copy Email",        command=self._copy_email)
        self.context_menu.add_command(label="🌐 Open Website",       command=self._open_website)
        self.context_menu.add_command(label="🗺️  Open Maps",         command=self._open_maps)
        self.tree.bind("<Button-3>", self._show_context_menu)
        self.tree.bind("<Double-1>", self._on_row_double_click)

    def _build_log_tab(self, parent):
        self.log_text = scrolledtext.ScrolledText(parent,
            font=("Consolas", 9), bg="#1e1e1e", fg="#d4d4d4",
            insertbackground="white", bd=0, padx=8, pady=8,
            state="disabled", wrap="word")
        self.log_text.pack(fill="both", expand=True)

        self.log_text.tag_configure("info",    foreground="#9cdcfe")
        self.log_text.tag_configure("success", foreground="#4ec9b0")
        self.log_text.tag_configure("warning", foreground="#dcdcaa")
        self.log_text.tag_configure("error",   foreground="#f44747")
        self.log_text.tag_configure("section", foreground="#569cd6",
            font=("Consolas", 9, "bold"))

        tk.Button(parent, text="Clear Log", font=("Segoe UI", 8),
            bg="#37474f", fg="white", bd=0, padx=6, pady=3,
            command=self._clear_log).pack(anchor="e", padx=4, pady=2)

    # ── Location cascade ──────────────────────────────────────────────────────
    def _on_country_change(self, *_):
        country = self.v_country.get()
        provinces = sorted(LOCATION_DATA.get(country, {}).keys())
        self.cb_province["values"] = provinces
        self.v_province.set(""); self.cb_city["values"] = []
        self.v_city.set(""); self.cb_area["values"] = []; self.v_area.set("")

    def _on_province_change(self, *_):
        country = self.v_country.get(); province = self.v_province.get()
        cities = sorted(LOCATION_DATA.get(country, {}).get(province, {}).keys())
        self.cb_city["values"] = cities
        self.v_city.set(""); self.cb_area["values"] = []; self.v_area.set("")

    def _on_city_change(self, *_):
        country = self.v_country.get(); province = self.v_province.get()
        city = self.v_city.get()
        areas = sorted(LOCATION_DATA.get(country, {}).get(province, {}).get(city, []))
        self.cb_area["values"] = areas; self.v_area.set("")

    # ── Scraping ──────────────────────────────────────────────────────────────
    def _validate_inputs(self):
        if not self.v_country.get():
            messagebox.showwarning("Missing Input", "Please select a Country."); return False
        if not self.v_city.get():
            messagebox.showwarning("Missing Input", "Please select a City."); return False
        category = self.v_custom_cat.get().strip() or self.v_category.get()
        if not category:
            messagebox.showwarning("Missing Input", "Please select or type a Business Category."); return False
        if not REQUESTS_OK or not BS4_OK:
            messagebox.showerror("Missing Libraries",
                "Required libraries not installed.\n\nRun:\n  pip install requests beautifulsoup4 lxml")
            return False
        return True

    def _start_scraping(self):
        if not self._validate_inputs(): return
        if self.is_scraping:
            messagebox.showinfo("Scraping in Progress", "Already scraping."); return

        self.is_scraping = True
        self.btn_search.config(state="disabled", bg="#9E9E9E")
        self.btn_stop.config(state="normal")
        self._set_status("🔄 Scraping…", "#FFF176")
        self.progress_var.set(0)
        self.lbl_progress.config(text="0%")

        country  = self.v_country.get()
        province = self.v_province.get()
        city     = self.v_city.get()
        area     = self.v_area.get().strip()
        category = self.v_custom_cat.get().strip() or self.v_category.get()
        max_res  = int(self.v_limit.get())

        loc_display = f"{area}, {city}, {province}, {country}".strip(", ")
        self.lbl_location_display.config(text=f"📍 {loc_display}  |  📦 {category}")
        self.lbl_count.config(text="  Total Found: 0  |  Showing: 0")

        self._log(f"{'='*50}", "section")
        self._log(f"  NEW SEARCH STARTED", "section")
        self._log(f"{'='*50}", "section")

        self.scrape_thread = threading.Thread(
            target=run_scraper,
            args=(country, province, city, area, category, max_res,
                  self._log_threadsafe, self._update_progress,
                  self._on_result_found, self._on_scraping_done),
            daemon=True
        )
        self.scrape_thread.start()

    def _stop_scraping(self):
        self.is_scraping = False
        self._set_status("⛔ Stopped", "#EF9A9A")
        self.btn_search.config(state="normal", bg=self.C_BTN)
        self.btn_stop.config(state="disabled")
        self._log("⛔ Scraping stopped by user.", "warning")

    def _log_threadsafe(self, msg):
        self.after(0, lambda m=msg: self._log(m))

    def _update_progress(self, value):
        self.after(0, lambda v=value: self._set_progress(v))

    def _on_result_found(self, biz):
        if self.is_scraping:
            self.after(0, lambda b=biz: self._add_row(b))

    def _on_scraping_done(self, businesses):
        self.businesses = businesses
        self.after(0, self._finish_scraping)

    def _finish_scraping(self):
        self.is_scraping = False
        self.btn_search.config(state="normal", bg=self.C_BTN)
        self.btn_stop.config(state="disabled")
        count = len(self.businesses)

        if count == 0:
            self._set_status("⚠️ No results found", "#FFE082")
            messagebox.showinfo("No Results",
                "No businesses found.\n\nTips:\n"
                "• Try a different area or city\n"
                "• Use a broader category\n"
                "• Check your internet connection\n"
                "• Google may require CAPTCHA verification")
        else:
            self._set_status(f"✅ Done – {count} businesses found", "#A5D6A7")

        self._refresh_count()

    def _add_row(self, biz):
        idx = len(self.tree.get_children())
        tag = "even" if idx % 2 == 0 else "odd"
        if biz.get("phone") or biz.get("email"):
            tag = "has_contact"
        elif "Google" in biz.get("source", ""):
            tag = "google"

        maps = "📍" if biz.get("maps_link") else ""
        web  = biz.get("website", "")[:50]

        self.tree.insert("", "end", values=(
            biz.get("name", ""),
            biz.get("category", ""),
            biz.get("address", ""),
            biz.get("phone", ""),
            biz.get("email", ""),
            web,
            biz.get("website_status", ""),
            biz.get("source", ""),
            maps,
        ), tags=(tag,))
        self._refresh_count()

    def _apply_filter(self, *_):
        query = self.v_filter.get().lower().strip()
        for item in self.tree.get_children():
            self.tree.delete(item)
        for idx, biz in enumerate(self.businesses):
            if not query or any(query in str(v).lower() for v in biz.values()):
                tag = "even" if idx % 2 == 0 else "odd"
                if biz.get("phone") or biz.get("email"):
                    tag = "has_contact"
                maps = "📍" if biz.get("maps_link") else ""
                self.tree.insert("", "end", values=(
                    biz.get("name", ""),
                    biz.get("category", ""),
                    biz.get("address", ""),
                    biz.get("phone", ""),
                    biz.get("email", ""),
                    biz.get("website", "")[:50],
                    biz.get("website_status", ""),
                    biz.get("source", ""),
                    maps,
                ), tags=(tag,))
        self._refresh_count()

    def _sort_column(self, col):
          data = [(self.tree.set(k, col), k) for k in self.tree.get_children("")]
        data.sort()
        for idx, (_, k) in enumerate(data):
            self.tree.move(k, "", idx)

    def _refresh_count(self):
        total   = len(self.businesses)
        showing = len(self.tree.get_children())
        self.lbl_count.config(text=f"  Total Found: {total}  |  Showing: {showing}")

    def _get_selected_biz(self):
        sel = self.tree.selection()
        if not sel: return None
        vals = self.tree.item(sel[0], "values")
        name = vals[0] if vals else ""
        for biz in self.businesses:
            if biz.get("name", "") == name:
                return biz
        return None

    def _show_context_menu(self, event):
        item = self.tree.identify_row(event.y)
        if item:
            self.tree.selection_set(item)
            self.context_menu.post(event.x_root, event.y_root)

    def _copy_name(self):
        biz = self._get_selected_biz()
        if biz: self.clipboard_clear(); self.clipboard_append(biz.get("name", ""))

    def _copy_phone(self):
        biz = self._get_selected_biz()
        if biz: self.clipboard_clear(); self.clipboard_append(biz.get("phone", ""))

    def _copy_email(self):
        biz = self._get_selected_biz()
        if biz: self.clipboard_clear(); self.clipboard_append(biz.get("email", ""))

    def _open_website(self):
        import webbrowser
        biz = self._get_selected_biz()
        if biz and biz.get("website"): webbrowser.open(biz["website"])

    def _open_maps(self):
        import webbrowser
        biz = self._get_selected_biz()
        if biz and biz.get("maps_link"): webbrowser.open(biz["maps_link"])

    def _on_row_double_click(self, event):
        self._open_website()

    def _search_info(self):
        country  = self.v_country.get(); province = self.v_province.get()
        city     = self.v_city.get();    area     = self.v_area.get()
        category = self.v_custom_cat.get().strip() or self.v_category.get()
        loc = f"{area}, {city}, {province}, {country}".strip(", ")
        return {"location": loc, "category": category}

    def _export_csv(self):
        if not self.businesses:
            messagebox.showwarning("No Data", "No results to export."); return
        fp = filedialog.asksaveasfilename(defaultextension=".csv",
            filetypes=[("CSV files", "*.csv")], title="Export as CSV")
        if fp:
            try:
                export_csv(self.businesses, fp)
                messagebox.showinfo("Exported", f"CSV saved:\n{fp}")
            except Exception as e:
                messagebox.showerror("Export Error", str(e))

    def _export_excel(self):
        if not self.businesses:
            messagebox.showwarning("No Data", "No results to export."); return
        fp = filedialog.asksaveasfilename(defaultextension=".xlsx",
            filetypes=[("Excel files", "*.xlsx")], title="Export as Excel")
        if fp:
            try:
                export_excel(self.businesses, fp, self._search_info())
                messagebox.showinfo("Exported", f"Excel saved:\n{fp}")
            except Exception as e:
                messagebox.showerror("Export Error", str(e))

    def _export_word(self):
        if not self.businesses:
            messagebox.showwarning("No Data", "No results to export."); return
        fp = filedialog.asksaveasfilename(defaultextension=".docx",
            filetypes=[("Word files", "*.docx")], title="Export as Word")
        if fp:
            try:
                export_word(self.businesses, fp, self._search_info())
                messagebox.showinfo("Exported", f"Word saved:\n{fp}")
            except Exception as e:
                messagebox.showerror("Export Error", str(e))

    def _clear_results(self):
        if messagebox.askyesno("Clear", "Clear all results?"):
            self.businesses.clear()
            for item in self.tree.get_children():
                self.tree.delete(item)
            self._refresh_count()
            self.lbl_location_display.config(text="")
            self._set_progress(0)

    def _set_status(self, text, bg="#a5d6a7"):
        self.status_label.config(text=f"● {text}", bg=self.C_SIDEBAR,
            fg=bg if bg.startswith("#") else "#a5d6a7")

    def _set_progress(self, value):
        self.progress_var.set(value)
        self.lbl_progress.config(text=f"{value}%")

    def _log(self, msg, tag="info"):
        self.log_text.config(state="normal")
        ts = datetime.now().strftime("%H:%M:%S")
        if "error" in msg.lower() or "⚠" in msg or "❌" in msg:
            tag = "error"
        elif "✅" in msg or "🎉" in msg or "complete" in msg.lower():
            tag = "success"
        elif "===" in msg or "───" in msg:
            tag = "section"
        self.log_text.insert("end", f"[{ts}] {msg}\n", tag)
        self.log_text.see("end")
        self.log_text.config(state="disabled")

    def _clear_log(self):
        self.log_text.config(state="normal")
        self.log_text.delete("1.0", "end")
        self.log_text.config(state="disabled")

    def _check_dependencies(self):
        missing = []
        if not REQUESTS_OK:  missing.append("requests")
        if not BS4_OK:       missing.append("beautifulsoup4")
        if not PANDAS_OK:    missing.append("pandas")
        if not OPENPYXL_OK:  missing.append("openpyxl  (Excel export)")
        if not DOCX_OK:      missing.append("python-docx  (Word export)")

        if missing:
            self._log("⚠  Some libraries missing:", "warning")
            self._log(f"   pip install {' '.join(m.split()[0] for m in missing)}", "warning")
        else:
            self._log(" All dependencies detected. Ready to scrape!", "success")
            self._log(" Sources: Google Local → Google Maps → Google → DDG → Bing", "info")
            self._log("Select a location & category then click 'Start Scraping'.", "info")


# ═══════════════════════════════════════════════════════════════════════════════
#  ENTRY POINT
# ═══════════════════════════════════════════════════════════════════════════════

if __name__ == "__main__":
    app = BusinessScraperApp()
    app.mainloop()